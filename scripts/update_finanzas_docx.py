"""Update financial tables/narrative via python-docx (no Word COM)."""
from copy import deepcopy
from pathlib import Path

from docx import Document

PATHS = [
    Path(r"C:\Users\Alisson.Munoz\Downloads\Nexus_Campus_Documentacion_actualizada.docx"),
    Path(r"C:\Users\Alisson.Munoz\Downloads\Nexus_Campus_Documentacion (1).docx"),
]

ROWS_VIAJES = [
    ("1-2", "320", "$0.215", "$69", "$0", "$69"),
    ("3", "960", "$0.215", "$206", "$0", "$206"),
    ("6", "2 880", "$0.215", "$619", "$500", "$1 119"),
    ("9", "4 480", "$0.215", "$963", "$500", "$1 463"),
    ("12", "5 600", "$0.215", "$1 204", "$500", "$1 704"),
]

ROWS_UTIL = [
    ("1-2", "69", "162", "-93"),
    ("3", "206", "162", "44"),
    ("6", "1 119", "162", "957"),
    ("9", "1 463", "162", "1 301"),
    ("12", "1 704", "162", "1 542"),
]

COST_ROWS = [
    ("Publicidad en redes sociales (Instagram, Facebook, TikTok)", "$100"),
    ("Infraestructura Appwrite / cloud (plan proyectado)", "$45"),
    ("Herramientas, dominio y operación menor", "$17"),
    ("Total mensual (costos fijos proyectados)", "$162"),
]


def set_para(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def replace_para_containing(doc, marker, new_text):
    for p in doc.paragraphs:
        if marker in p.text:
            set_para(p, new_text)
            return True
    return False


def replace_substr(doc, old, new):
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            set_para(p, p.text.replace(old, new))
            n += 1
    return n


def clear_table_body(table):
    # remove rows from bottom except header
    tbl = table._tbl
    trs = list(tbl.tr_lst)
    for tr in trs[1:]:
        tbl.remove(tr)


def fill_rows(table, rows):
    for values in rows:
        row = table.add_row()
        for i, val in enumerate(values):
            row.cells[i].text = val


def update_tables(doc):
    for table in doc.tables:
        h0 = table.rows[0].cells[0].text.strip().lower()
        h1 = table.rows[0].cells[1].text.strip().lower() if len(table.rows[0].cells) > 1 else ""
        h2 = table.rows[0].cells[2].text.strip().lower() if len(table.rows[0].cells) > 2 else ""

        if h0.startswith("actividad") and "costo" in h1:
            clear_table_body(table)
            fill_rows(table, COST_ROWS)
            print("  cost table OK")
        elif h0 == "mes" and "viaje" in h1:
            clear_table_body(table)
            fill_rows(table, ROWS_VIAJES)
            print("  viajes/ingresos table OK")
        elif h0 == "mes" and "ingreso" in h1 and "costo" in h2:
            clear_table_body(table)
            fill_rows(table, ROWS_UTIL)
            print("  utilidad table OK")


def update_doc(path: Path) -> bool:
    if not path.exists():
        print("SKIP missing", path)
        return False
    print("Updating", path)
    try:
        doc = Document(str(path))
    except Exception as e:
        print("  open fail:", e)
        return False

    replace_para_containing(
        doc,
        "Para efectos de la proyección financiera, se asumió una distancia promedio de 3 km",
        "Para efectos de la proyección financiera, se asumió una distancia promedio de 3 km "
        "por viaje, obteniendo una tarifa promedio de $2.15 (coherente con la fórmula "
        "implementada en la app: $0.80 + $0.45/km, mínimo $1.00). El plan de negocio proyecta "
        "una comisión del 10 % sobre el valor de cada viaje ($0.215 en el escenario promedio). "
        "Esta comisión es del modelo financiero y no se cobra automáticamente en la versión "
        "actual (1.2.13). Los viajes mensuales se estiman como: Viajes = Usuarios activos × 8 "
        "/ 3, porque cada viaje involucra en promedio un conductor y dos pasajeros; así se "
        "evita contar tres veces el mismo viaje cuando varios usuarios activos participan.",
    )
    replace_para_containing(
        doc,
        "Este valor resulta sostenible al compararlo con los ingresos que genera un usuario activo",
        "Este valor resulta sostenible al compararlo con el aporte económico de un usuario activo. "
        "Si cada usuario participa en promedio en 8 viajes al mes y cada viaje involucra ~3 "
        "participantes, el aporte mensual por usuario vía comisión es aproximadamente "
        "(8/3) × $0.215 ≈ $0.57. En 6 meses el LTV aproximado es ~$3.4, superior al CAC de "
        "$0.34. El MRR proyectado en el mes 12 (comisión + convenio) alcanza ~$1.704.",
    )
    replace_para_containing(
        doc,
        "Los ingresos proyectados de Nexus Campus provienen de dos fuentes: la comisión",
        "Los ingresos proyectados de Nexus Campus provienen de dos fuentes: la comisión "
        "estimada por cada viaje realizado (modelo financiero, no cobrada aún en la app) "
        "y los convenios institucionales con universidades. El número de viajes mensuales "
        "se estima con Viajes = Usuarios activos × 8 / 3. El ingreso por comisión es "
        "Viajes × $0.215. A partir del mes 6 se incorpora un convenio institucional fijo "
        "de $500 mensuales.",
    )
    replace_para_containing(
        doc,
        "El análisis muestra que Nexus Campus logra cubrir sus costos operativos desde los primeros meses",
        "Con costos fijos proyectados de $162 mensuales (marketing $100 + infraestructura $45 + "
        "operación $17), el punto de equilibrio se alcanza cerca de 753 viajes/mes "
        "(≈ 94 usuarios activos). En la proyección corregida, el primer bimestre aún presenta "
        "pérdida operativa ($-93), el mes 3 alcanza utilidad positiva (~$44) y desde el mes 6, "
        "con el convenio institucional, la utilidad se fortalece de forma consistente.",
    )

    # entregables note
    for marker in (
        "Finalmente, el documento puede complementarse con material audiovisual",
        "Finalmente, el documento incluye el material audiovisual",
        "Los entregables audiovisuales del examen se complementan así",
    ):
        if replace_para_containing(
            doc,
            marker,
            "Los entregables audiovisuales del examen se complementan así: el pitch deck está "
            "en docs/Nexus_Campus_Pitch_Deck.pptx del repositorio; el video promocional se "
            "anexa mediante el enlace oficial del equipo (PENDIENTE: pegar link del video). "
            "El README del repositorio describe instalación, configuración y ejecución. "
            "La versión documentada de la app es Nexus Campus 1.2.13.",
        ):
            break

    replace_substr(
        doc,
        "se estableció un presupuesto mensual fijo de $100, destinado a la ejecución de campañas",
        "se estableció un presupuesto de adquisición de $100 mensuales dentro de una estructura "
        "de costos fijos de $162 (marketing $100, infraestructura proyectada $45 y operación $17), "
        "destinado a campañas",
    )

    update_tables(doc)

    out = path
    try:
        doc.save(str(out))
        print("  SAVED", out)
    except PermissionError:
        alt = path.with_name(path.stem + "_finanzas_ok.docx")
        doc.save(str(alt))
        print("  LOCKED; saved", alt)
    return True


def main():
    for p in PATHS:
        update_doc(p)


if __name__ == "__main__":
    main()
