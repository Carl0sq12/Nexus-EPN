from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
import shutil

src = r"C:\Users\Alisson.Munoz\Downloads\Nexus_Campus_Documentacion (1).docx"
backup = r"C:\Users\Alisson.Munoz\Downloads\Nexus_Campus_Documentacion_backup.docx"
out = r"C:\Users\Alisson.Munoz\Downloads\Nexus_Campus_Documentacion_actualizada.docx"

try:
    shutil.copy2(src, backup)
except PermissionError:
    print("WARN: no se pudo crear backup (archivo abierto). Continuando...")

doc = Document(src)


def set_para_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_in_paragraphs(contains: str, new_text: str, count: int = 1) -> int:
    n = 0
    for p in doc.paragraphs:
        if contains in p.text:
            set_para_text(p, new_text)
            n += 1
            if n >= count:
                break
    return n


def replace_substring(old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            set_para_text(p, p.text.replace(old, new))
            n += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_para_text(p, p.text.replace(old, new))
                        n += 1
    return n


# 1) Intro audiovisual + version
replace_in_paragraphs(
    "Finalmente, el documento incluye el material audiovisual",
    "Finalmente, el documento puede complementarse con material audiovisual del proyecto "
    "(video promocional y pitch deck) cuando dicho material se incorpore como anexo a la "
    "sustentación. La versión actual de la aplicación móvil documentada corresponde a "
    "Nexus Campus 1.2.13.",
)

# 2) Comisión = proyección, no cobro en app
replace_in_paragraphs(
    "Comisión mínima por viaje:",
    "Comisión mínima por viaje (proyección de negocio): porcentaje proyectado sobre cada "
    "viaje compartido para cubrir costos operativos. En la versión actual del prototipo "
    "esta comisión no se cobra automáticamente dentro de la aplicación; el precio por "
    "asiento se calcula y negocia en la app, pero no existe pasarela de pagos ni retención "
    "automática de comisión.",
)

# 3) Modelo de precios / comisión proyectada
replace_in_paragraphs(
    "Para efectos de la proyección financiera, se asumió una distancia promedio de 3 km",
    "Para efectos de la proyección financiera, se asumió una distancia promedio de 3 km "
    "por viaje, obteniendo una tarifa promedio de $2.15. Asimismo, el plan de negocio "
    "proyecta una comisión del 10 % sobre el valor de cada viaje (equivalente a $0.215 "
    "por viaje en el escenario promedio). Esta comisión forma parte del modelo financiero "
    "proyectado y no está implementada como cobro automático en el código de la aplicación "
    "en su versión actual (1.2.13).",
)

# 4) Canales / APK
replace_in_paragraphs(
    "Tiendas de aplicaciones (Google Play / App Store):",
    "Tiendas de aplicaciones (Google Play / App Store): canal de distribución proyectado "
    "para el lanzamiento oficial. En la fase actual de prototipo académico la distribución "
    "se realiza principalmente mediante instalable Android (APK).",
)

# 5) nueve -> diez colecciones
replace_substring(
    "compuesta por nueve colecciones principales",
    "compuesta por diez colecciones principales",
)

# 6) Buckets
replace_in_paragraphs(
    "Adicionalmente, el proyecto utiliza dos buckets de almacenamiento",
    "Adicionalmente, el proyecto utiliza almacenamiento de archivos en Appwrite. El bucket "
    "avatars almacena fotografías de perfil de los usuarios. Las imágenes de vehículos y "
    "licencia se gestionan también sobre Appwrite; en el plan gratuito, cuando solo se "
    "dispone de un bucket, se reutiliza avatars con rutas del tipo vehicles/{id}.jpg. "
    "El tamaño máximo configurado es de 10 MB por archivo.",
)

# 7) Referencia de sección incorrecta
replace_in_paragraphs(
    "descritas en la sección 6.2.4",
    "descritas en la sección 2.4 (Manual de Despliegue), de acuerdo con lo solicitado en la "
    "guía de la actividad.",
)

# 8) Ingresos proyectados
replace_in_paragraphs(
    "Los ingresos de Nexus Campus provienen de dos fuentes: la comisión cobrada",
    "Los ingresos proyectados de Nexus Campus provienen de dos fuentes: la comisión "
    "estimada por cada viaje realizado (modelo financiero, no cobrada aún en la app) "
    "y los convenios institucionales con universidades. El número de viajes mensuales "
    "se estima multiplicando los usuarios activos por un promedio de 8 viajes al mes "
    "(2 viajes por semana × 4 semanas), y el ingreso por comisión se obtiene "
    "multiplicando los viajes mensuales por la comisión proyectada de $0.215 por viaje. "
    "A partir del mes 6 se incorpora un convenio institucional fijo de $500 mensuales "
    "con una universidad.",
)

# 9) Capacidades implementadas
for p in doc.paragraphs:
    if p.text.startswith(
        "Esta organización permite separar las responsabilidades"
    ):
        set_para_text(
            p,
            p.text
            + " Entre las capacidades implementadas en la versión 1.2.13 destacan: "
            "separación entre Solicitudes (cupos) y Notificaciones (chat, fin de viaje, SOS); "
            "solicitud de cupo con parada obligatoria sobre la ruta del conductor; "
            "navegación en tiempo real con recorte de la polilínea recorrida; "
            "finalización automática al llegar al destino; calificación al llegar a la parada "
            "del pasajero; aprobación de vehículo del conductor; y autenticación biométrica "
            "para desbloquear una sesión existente.",
        )
        break

# 10) BaaS wording
replace_substring(
    "equivalente funcional a un backend as a service tipo Supabase",
    "como Backend as a Service (BaaS)",
)

# 11) Realtime + trip_locations
replace_substring(
    "Suscripción WebSocket a las colecciones trips y notifications",
    "Suscripción WebSocket a colecciones como trips, notifications y trip_locations",
)

# 12) Add trip_locations row
for table in doc.tables:
    header = table.rows[0].cells[0].text.strip().lower()
    if not header.startswith("colecci"):
        continue
    existing = " ".join(r.cells[0].text for r in table.rows)
    if "trip_locations" not in existing:
        row = table.add_row()
        row.cells[0].text = "trip_locations"
        row.cells[1].text = (
            "Ubicación en vivo del conductor durante la navegación del viaje "
            "(seguimiento GPS para pasajeros)."
        )
    break

save_path = src
try:
    doc.save(src)
except PermissionError:
    save_path = out
    doc.save(out)
    print("WARN: original bloqueado; guardado en:", out)

print("Saved:", save_path)
print("Backup:", backup)

doc2 = Document(save_path)
text = "\n".join(p.text for p in doc2.paragraphs)
text += "\n".join(
    c.text for t in doc2.tables for r in t.rows for c in r.cells
)
checks = [
    "trip_locations",
    "diez colecciones",
    "no está implementada como cobro automático",
    "1.2.13",
    "APK",
    "sección 2.4",
    "vehicles/{id}.jpg",
    "no cobrada aún en la app",
]
for c in checks:
    print(("OK" if c in text else "MISSING"), "-", c)
